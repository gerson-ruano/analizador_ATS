import PyPDF2
from docx import Document
import streamlit as st
import re
import tempfile
import os
from typing import Optional, Tuple

try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text_from_uploaded_file(self, uploaded_file) -> Tuple[str, bool]:
        """Procesa archivos subidos a Streamlit y retorna texto y si fue exitoso"""
        try:
            # Crear archivo temporal
            with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.name) as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                temp_path = tmp_file.name
            
            try:
                # Extraer texto basado en el tipo de archivo
                if uploaded_file.type == "application/pdf":
                    text = self._extract_from_pdf(temp_path)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    text = self._extract_from_docx(temp_path)
                elif uploaded_file.type == "text/plain":
                    text = self._extract_from_txt(temp_path)
                else:
                    return f"Tipo de archivo no soportado: {uploaded_file.type}", False
                
                # Verificar si se extrajo texto válido
                if text and len(text.strip()) > 50:  # Mínimo 50 caracteres
                    return text, True
                else:
                    return "No se pudo extraer texto suficiente del documento", False
                    
            finally:
                # Limpiar archivo temporal
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            return f"Error procesando archivo: {str(e)}", False
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extrae texto de archivos PDF con manejo mejorado de errores"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                if len(reader.pages) == 0:
                    return "PDF vacío o corrupto"
                
                text = ""
                pages_with_text = 0
                
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"--- Página {i+1} ---\n{page_text}\n\n"
                        pages_with_text += 1
                
                if pages_with_text == 0 and OCR_AVAILABLE:
                    st.info("📄 PDF parece ser escaneado. Usando OCR...")
                    return self._extract_with_ocr(file_path)
                
                return text if text.strip() else "No se pudo extraer texto del PDF"
                
        except Exception as e:
            return f"Error procesando PDF: {str(e)}"
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extrae texto de archivos Word con formato mejorado"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts) if text_parts else "Documento DOCX vacío"
            
        except Exception as e:
            return f"Error procesando DOCX: {str(e)}"
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extrae texto de archivos TXT con múltiples codificaciones"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'windows-1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    if content.strip():
                        return content
            except UnicodeDecodeError:
                continue
        
        return "No se pudo decodificar el archivo de texto con ninguna codificación común"
    
    def _extract_with_ocr(self, file_path: str) -> str:
        """Usa OCR para PDFs escaneados con manejo de errores"""
        if not OCR_AVAILABLE:
            return "OCR no disponible. Instala pdf2image y pytesseract para esta funcionalidad."
        
        try:
            images = convert_from_path(file_path, dpi=300)
            text = ""
            
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image, lang='spa')
                if page_text.strip():
                    text += f"--- Página {i+1} (OCR) ---\n{page_text}\n\n"
            
            return text if text.strip() else "OCR no pudo extraer texto"
            
        except Exception as e:
            return f"Error en OCR: {str(e)}"
    
    def get_document_stats(self, text: str) -> dict:
        """Obtiene estadísticas del documento"""
        if not text or text.startswith("Error") or text.startswith("No se pudo"):
            return {}
        
        words = text.split()
        lines = text.split('\n')
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        return {
            'caracteres': len(text),
            'palabras': len(words),
            'lineas': len(lines),
            'parrafos': len(paragraphs),
            'densidad_palabras': len(words) / max(len(lines), 1)
        }